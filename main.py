#安装pip install ttkbootstrap pandas openpyxl python-docx tkinterdnd2
# 学号抽取器 v2.0（混合数据版）
# 安装依赖：pip install python-docx  # 新增Word文档解析依赖
import ctypes
import sys
import os
import tkinter as tk
from tkinter import messagebox, filedialog
import random
import pandas as pd
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import ttkbootstrap as tbs
from ttkbootstrap.constants import *
from tkinterdnd2 import DND_FILES, TkinterDnD

# 系统级设置
try:
    ctypes.windll.shcore.SetProcessDpiAwareness(1)
except AttributeError:
    print("警告：无法设置 DPI Awareness，可能影响界面显示。")
sys.stdout = open(os.devnull, 'w') if 'ONEFILE' in os.environ else sys.stdout

class NumberPickerApp:
    def __init__(self, root):
        self.root = root
        self.language = "zh"  # 新增：默认语言
        self.translations = {
            "zh": {
                "title": "🎓 智能抽取器 v2.0",
                "import": "📂 导入数据文件",
                "not_loaded": "未加载文件",
                "setting": "⚙️ 设置",
                "range": "范围设置（示例：0-60）",
                "start": "起始编号:",
                "end": "结束编号:",
                "pick": "🚀 开始抽取",
                "reset": "🗑 重置数据",
                "click_to_start": "点击开始抽取",
                "scrolling": "🎚 随机滚动中...",
                "file_loaded": "文件已加载",
                "no_valid_data": "文件中无有效数据",
                "error": "错误",
                "confirm": "确认",
                "clear_all": "清除所有数据？",
                "font_size": "结果字体大小:",
                "pick_count": "一次抽取人数:",
                "speed": "抽取动画速度(ms):",
                "save": "保存",
                "settings": "设置",
                "language": "语言选择:",
                "chinese": "简体中文",
                "english": "English",
                "import_failed": "导入失败",
                "invalid_word": "❌ 无效的Word文件或文件已损坏",
                "invalid_setting": "设置无效",
                "range_too_large": "范围过大（{count}条），继续？",
                "end_gt_start": "结束编号必须大于起始编号",
                "must_int": "编号必须为整数（示例：0-60）",
                "no_candidates": "无有效候选数据",
            },
            "en": {
                "title": "🎓 SmartPicker v2.0",
                "import": "📂 Import Data File",
                "not_loaded": "No file loaded",
                "setting": "⚙️ Settings",
                "range": "Range Setting (e.g. 0-60)",
                "start": "Start No.:",
                "end": "End No.:",
                "pick": "🚀 Start Picking",
                "reset": "🗑 Reset Data",
                "click_to_start": "Click to Start",
                "scrolling": "🎚 Rolling...",
                "file_loaded": "File Loaded",
                "no_valid_data": "No valid data in file",
                "clear_all": "Clear all data?",
                "font_size": "Result Font Size:",
                "pick_count": "Pick Count:",
                "speed": "Animation Speed (ms):",
                "save": "Save",
                "settings": "Settings",
                "language": "Language:",
                "chinese": "简体中文",
                "english": "English",
                "import_failed": "Import Failed",
                "invalid_word": "❌ Invalid or corrupted Word file",
                "invalid_setting": "Invalid setting",
                "range_too_large": "Range too large ({count} items), continue?",
                "end_gt_start": "End number must be greater than start number",
                "must_int": "Numbers must be integers (e.g. 0-60)",
                "no_candidates": "No valid candidates",
            }
        }
        self.root.title(self.t("title"))
        self.root.geometry("920x680")

        # 初始化拖放功能
        self.root.drop_target_register(DND_FILES)
        self.root.dnd_bind('<<Drop>>', self.handle_file_drop)
        
        # 核心配置
        self.DEFAULT_FONT = ("微软雅黑", 14)
        self.MAX_RANGE = 50000
        self.ANIMATION_FRAMES = 60
        self.DEFAULT_START = "0"
        self.DEFAULT_END = "60"
        self.PICK_COUNT = 1  # 新增：一次抽取人数
        self.FONT_SIZE = 30  # 新增：结果字体大小
        self.ANIMATION_SPEED = 20  # 新增：动画速度(ms)
        self.allow_duplicates = False  # 新增：是否允许重复抽取
        
        # 数据存储
        self.excel_path = None
        self.raw_data = []
        self.clean_data = []
        
        # 初始化界面
        self.create_widgets()
        self.configure_styles()
        self.reset_ui()

    def t(self, key, **kwargs):
        """多语言文本获取"""
        text = self.translations[self.language].get(key, key)
        return text.format(**kwargs) if kwargs else text

    def create_widgets(self):
        main_frame = tbs.Frame(self.root, padding=20)
        main_frame.pack(fill=BOTH, expand=YES, padx=15, pady=15)

        # 文件操作区
        file_frame = tbs.Frame(main_frame, padding=10)
        file_frame.pack(fill=X, pady=10)
        
        self.import_btn = tbs.Button(file_frame, 
                                   text=self.t("import"), 
                                   command=self.import_excel,
                                   bootstyle=(SUCCESS, OUTLINE, "rounded"))
        self.import_btn.pack(side=LEFT, padx=5)
        
        self.file_status = tbs.Label(file_frame, 
                                   text=self.t("not_loaded"), 
                                   font=self.DEFAULT_FONT,
                                   foreground="#7f8c8d")
        self.file_status.pack(side=LEFT, padx=15)

        # 新增设置按钮
        self.setting_btn = tbs.Button(file_frame,
                                   text=self.t("setting"),
                                   command=self.open_settings,
                                   bootstyle=(INFO, OUTLINE, "rounded"))
        self.setting_btn.pack(side=LEFT, padx=5)

        # 输入设置区
        input_frame = tbs.Labelframe(main_frame, text=self.t("range"), 
                                   padding=15, bootstyle=SECONDARY)
        input_frame.pack(fill=X, pady=10, padx=5)
        
        tbs.Label(input_frame, text=self.t("start")).grid(row=0, column=0, sticky=W, pady=3)
        self.start_entry = tbs.Entry(input_frame, width=8, font=self.DEFAULT_FONT)
        self.start_entry.grid(row=0, column=1, padx=5, pady=3)
        self.start_entry.insert(0, self.DEFAULT_START)
        
        tbs.Label(input_frame, text=self.t("end")).grid(row=1, column=0, sticky=W, pady=3)
        self.end_entry = tbs.Entry(input_frame, width=8, font=self.DEFAULT_FONT)
        self.end_entry.grid(row=1, column=1, padx=5, pady=3)
        self.end_entry.insert(0, self.DEFAULT_END)

        # 结果显示区
        self.result_frame = tbs.Frame(main_frame, padding=30, relief="sunken", bootstyle="dark")
        self.result_frame.pack(fill=BOTH, expand=YES, pady=20, padx=5)
        # 新增居中容器
        center_container = tbs.Frame(self.result_frame)
        center_container.pack(expand=True, fill=BOTH, anchor='center')

        self.result_label = tbs.Label(center_container, 
                                text=self.t("click_to_start"), 
                                style="Result.TLabel",
                                wraplength=1500, #换行距离
                                font=("微软雅黑", self.FONT_SIZE)) # 使用可变字体大小
        self.result_label.pack(expand=True, anchor='center')  # 修改pack参数
        # 控制按钮
        control_frame = tbs.Frame(main_frame, padding=10)
        control_frame.pack(fill=X, pady=15)
        
        self.start_btn = tbs.Button(control_frame, 
                                  text=self.t("pick"), 
                                  command=self.start_pick,
                                  bootstyle=(SUCCESS, "outline-toolbutton", "inverse"),
                                  style="LargeButton.TButton",
                                  padding=12)
        self.start_btn.pack(side=LEFT, padx=10)
        
        self.reset_btn = tbs.Button(control_frame, 
                                  text=self.t("reset"), 
                                  command=self.reset_data,
                                  bootstyle=(DANGER, "outline-toolbutton"))
        self.reset_btn.pack(side=LEFT, padx=10)

    def configure_styles(self):
        self.style = tbs.Style()
        self.style.configure(
            "Result.TLabel",
            font=("微软雅黑", 80, "bold"),
            foreground="#2ecc71",
            borderwidth=6,
            padding=60,
            anchor='center'
        )  # 确保括号正确闭合
    def reset_ui(self):
        self.start_entry.delete(0, tk.END)
        self.start_entry.insert(0, self.DEFAULT_START)
        self.end_entry.delete(0, tk.END)
        self.end_entry.insert(0, self.DEFAULT_END)
        self.file_status.config(text=self.t("not_loaded"), foreground="#7f8c8d")
        self.result_label.config(text=self.t("click_to_start"), foreground="#2ecc71")
        self.raw_data = []
        self.clean_data = []

    def import_excel(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("数据文件", "*.xlsx;*.xls;*.csv;*.txt;*.docx")]  # 添加docx,excel,csv,txt,docx格式
        )
        if not file_path: return

        self.import_file(file_path)

    def import_file(self, file_path):
        """导入文件的逻辑"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext in ('.xlsx', '.xls'):
                # 处理Excel文件
                df = pd.read_excel(file_path, engine="openpyxl", dtype=str, nrows=200)
                valid_cols = self.detect_valid_columns(df)

                if not valid_cols:
                    raise ValueError("未发现有效数据列（至少2个非空项）")

                df = pd.read_excel(file_path, engine="openpyxl", usecols=valid_cols[:1], dtype=str)
                self.raw_data = df.iloc[:, 0].dropna().astype(str).tolist()
                
            elif file_ext == '.csv':
                # 处理CSV文件
                df = pd.read_csv(file_path, dtype=str, nrows=200, encoding='utf-8', sep=None, engine='python')
                valid_cols = self.detect_valid_columns(df)

                if not valid_cols:
                    raise ValueError("未发现有效数据列（至少2个非空项）")

                self.raw_data = df.iloc[:, 0].dropna().astype(str).tolist()
                
            elif file_ext == '.txt':
                # 处理TXT文件
                with open(file_path, 'r', encoding='utf-8') as f:
                    self.raw_data = [line.strip() for line in f.readlines()]
            elif file_ext == '.docx':
                try:
                    doc = Document(file_path)
                    self.raw_data = []
                    for table in doc.tables:
                        for row in table.rows:
                            row_data = [cell.text.strip() for cell in row.cells]
                            self.raw_data.extend([d for d in row_data if d])
                    if not self.raw_data:
                        for para in doc.paragraphs:
                            text = para.text.strip()
                            if text:
                                self.raw_data.append(text)
                    self.clean_data = [item.replace('\n', ' ').strip() for item in self.raw_data if item.strip()]
                except PackageNotFoundError:
                    raise ValueError(self.t("invalid_word"))
                except Exception as e:
                    messagebox.showerror(self.t("import_failed"), f"❌ {self.t('error')}: {str(e)}")
                    self._reset_import_state()
            # 通用数据清洗
            self.clean_data = [item.strip() for item in self.raw_data if item.strip() != '']

            if not self.clean_data:
                raise ValueError(self.t("no_valid_data"))

            self.excel_path = file_path
            self.file_status.config(
                text=f"📂 {os.path.basename(file_path)} ({len(self.clean_data)}条)",
                foreground="#2ecc71"
            )
            self.root.after(3000, lambda: self.file_status.config(
                text=self.t("file_loaded"), foreground="#7f8c8d"
            ))
            print(f"导入文件路径: {file_path}")
            print(f"清洗后的数据: {self.clean_data}")

        except Exception as e:
            messagebox.showerror(self.t("import_failed"), f"❌ {self.t('error')}: {str(e)}")
            self._reset_import_state()

    def detect_valid_columns(self, df):
        """检测包含有效数据的列（至少2个非空项）"""
        valid_cols = []
        for col in df.columns:
            sample = df[col].dropna().astype(str).str.strip()
            if len(sample) >= 2:
                valid_cols.append(col)
        return valid_cols

    def _reset_import_state(self):
        self.excel_path = None
        self.raw_data = []
        self.clean_data = []
        self.file_status.config(text=self.t("not_loaded"), foreground="#7f8c8d")

    def validate_input(self):
        if self.clean_data:
            return True

        try:
            start = int(self.start_entry.get())
            end = int(self.end_entry.get())
            
            if end - start > self.MAX_RANGE:
                if not messagebox.askyesno(self.t("confirm"), self.t("range_too_large", count=end-start+1)):
                    return False
                    
            if start > end:
                raise ValueError(self.t("end_gt_start"))
                
            return True

        except ValueError:
            messagebox.showerror(self.t("error"), self.t("must_int"))
            return False

    def start_pick(self):
        if not self.validate_input():
            self.start_btn.config(state=NORMAL)
            return

        self.start_btn.config(state=DISABLED)
        self.result_label.config(text=self.t("scrolling"), foreground="#f1c40f")
        
        # 生成候选数据
        candidates = self.clean_data if self.clean_data else self.generate_manual_candidates()
        
        if not candidates:
            messagebox.showwarning(self.t("error"), self.t("no_candidates"))
            self.start_btn.config(state=NORMAL)
            return

        # 打乱数据
        random.shuffle(candidates)
        self.shuffled_data = candidates
        
        self.animation_step = 0
        self.run_animation()

    def generate_manual_candidates(self):
        try:
            start = int(self.start_entry.get())
            end = int(self.end_entry.get())
            return [str(num) for num in range(start, end+1)]
        except:
            return []

    def run_animation(self):
        if self.animation_step < self.ANIMATION_FRAMES:
            # 支持多人数抽取
            pick_count = min(self.PICK_COUNT, len(self.shuffled_data))
            if self.allow_duplicates:
                # 允许重复抽取
                result = "\n".join(random.choices(self.shuffled_data, k=pick_count))
            else:
                # 不允许重复抽取
                indices = random.sample(range(len(self.shuffled_data)), pick_count)
                result = "\n".join([self.shuffled_data[i] for i in indices])
            self.result_label.config(
                text=result,
                foreground="#e74c3c" if self.animation_step % 3 else "#3498db"
            )
            self.animation_step += 1
            self.root.after(self.ANIMATION_SPEED, self.run_animation)
        else:
            pick_count = min(self.PICK_COUNT, len(self.shuffled_data))
            if self.allow_duplicates:
                result = "\n".join(random.choices(self.shuffled_data, k=pick_count))
            else:
                result = "\n".join(self.shuffled_data[:pick_count])
            self.result_label.config(text=result, foreground="#2ecc71")
            self.start_btn.config(state=NORMAL)
            self.root.bell()

    def reset_data(self):
        if messagebox.askyesno(self.t("confirm"), self.t("clear_all")):
            self._reset_import_state()
            self.reset_ui()

    # 新增：设置窗口
    def open_settings(self):
        setting_win = tbs.Toplevel(self.root)
        setting_win.title(self.t("settings"))
        setting_win.geometry("400x500")  # 调整窗口高度
        setting_win.resizable(False, False)

        tbs.Label(setting_win, text=self.t("font_size"), font=self.DEFAULT_FONT).pack(pady=10)
        font_size_var = tk.IntVar(value=self.FONT_SIZE)
        font_size_entry = tbs.Entry(setting_win, textvariable=font_size_var, width=8)
        font_size_entry.pack()

        tbs.Label(setting_win, text=self.t("pick_count"), font=self.DEFAULT_FONT).pack(pady=10)
        pick_count_var = tk.IntVar(value=self.PICK_COUNT)
        pick_count_entry = tbs.Entry(setting_win, textvariable=pick_count_var, width=8)
        pick_count_entry.pack()

        tbs.Label(setting_win, text=self.t("speed"), font=self.DEFAULT_FONT).pack(pady=10)
        speed_var = tk.IntVar(value=self.ANIMATION_SPEED)
        speed_entry = tbs.Entry(setting_win, textvariable=speed_var, width=8)
        speed_entry.pack()

        # 新增语言选择
        tbs.Label(setting_win, text=self.t("language"), font=self.DEFAULT_FONT).pack(pady=10)
        lang_var = tk.StringVar(value=self.language)
        lang_frame = tbs.Frame(setting_win)
        lang_frame.pack()
        tbs.Radiobutton(lang_frame, text=self.t("chinese"), variable=lang_var, value="zh").pack(side=LEFT, padx=10)
        tbs.Radiobutton(lang_frame, text=self.t("english"), variable=lang_var, value="en").pack(side=LEFT, padx=10)

        # 新增“是否允许重复抽取”选项
        allow_duplicates_var = tk.BooleanVar(value=self.allow_duplicates)
        tbs.Label(setting_win, text="是否允许重复抽取:", font=self.DEFAULT_FONT).pack(pady=10)
        allow_duplicates_check = tbs.Checkbutton(
            setting_win, 
            text="允许重复", 
            variable=allow_duplicates_var, 
            bootstyle="round-toggle"
        )
        allow_duplicates_check.pack()

        def save_settings():
            try:
                self.FONT_SIZE = max(10, int(font_size_var.get()))
                self.PICK_COUNT = max(1, int(pick_count_var.get()))
                self.ANIMATION_SPEED = max(5, int(speed_var.get()))
                self.language = lang_var.get()
                self.allow_duplicates = allow_duplicates_var.get()  # 保存“是否允许重复抽取”选项
                self.result_label.config(font=("微软雅黑", self.FONT_SIZE))
                # 重新刷新界面文本
                self.refresh_texts()
                setting_win.destroy()
            except Exception as e:
                messagebox.showerror(self.t("error"), f"{self.t('invalid_setting')}: {e}")

        tbs.Button(setting_win, text=self.t("save"), command=save_settings, bootstyle=SUCCESS).pack(pady=18)

    # 新增：刷新界面所有文本
    def refresh_texts(self):
        self.root.title(self.t("title"))
        self.import_btn.config(text=self.t("import"))
        self.file_status.config(text=self.t("not_loaded"))
        self.setting_btn.config(text=self.t("setting"))
        self.result_label.config(text=self.t("click_to_start"))
        self.start_btn.config(text=self.t("pick"))  # 刷新“开始抽取”按钮
        self.reset_btn.config(text=self.t("reset"))  # 刷新“重置数据”按钮
        # 这里只刷新主要按钮和标签，复杂布局可重启程序生效

    def handle_file_drop(self, event):
        file_path = event.data.strip().strip('{}')  # 移除路径中的花括号
        if os.path.isfile(file_path):
            self.import_file(file_path)

if __name__ == "__main__":
    # 使用 TkinterDnD 的窗口类
    root = TkinterDnD.Tk()  
    # 手动应用 ttkbootstrap 主题
    style = tbs.Style(theme="darkly")  
    root.geometry("920x680")
    app = NumberPickerApp(root)
    root.mainloop()